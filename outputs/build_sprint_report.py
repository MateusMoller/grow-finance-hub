from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from datetime import date

OUT = r"D:\grow-finance-hub-main\outputs\Relatorio_Sprint_Grow_Finance_Hub_06-08-2026.docx"
NAVY = "0B2545"
BLUE = "2E74B5"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
MUTED = "667085"
GREEN = "16794B"
AMBER = "8A6100"
RED = "A11B1B"

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, NAVY, 8, 4),
]:
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_name in ["List Bullet", "List Number"]:
    style = styles[list_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(11)
    style.paragraph_format.left_indent = Inches(0.5)
    style.paragraph_format.first_line_indent = Inches(-0.25)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.167

if "Lead" not in styles:
    lead = styles.add_style("Lead", WD_STYLE_TYPE.PARAGRAPH)
else:
    lead = styles["Lead"]
lead.font.name = "Calibri"
lead.font.size = Pt(12)
lead.font.color.rgb = RGBColor.from_string(NAVY)
lead.paragraph_format.space_after = Pt(10)
lead.paragraph_format.line_spacing = 1.15

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)

def setup_header_footer(sec):
    hp = sec.header.paragraphs[0]
    hp.text = "GROW FINANCE HUB  |  RELATÓRIO DA SPRINT"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.runs[0]
    hr.font.name = "Calibri"
    hr.font.size = Pt(8.5)
    hr.font.bold = True
    hr.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_number(sec.footer.paragraphs[0])

setup_header_footer(section)

def add_title(text, size, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    return p

def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

def add_numbered(items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)

def add_status_table(rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Frente", "Situação", "Observação"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, PALE_BLUE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(NAVY)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for front, status, note in rows:
        cells = table.add_row().cells
        for idx, value in enumerate((front, status, note)):
            cells[idx].text = value
            cells[idx].paragraphs[0].paragraph_format.space_after = Pt(0)
        color = GREEN if status == "Implementado" else AMBER if status == "Pendente" else BLUE
        cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(color)
        cells[1].paragraphs[0].runs[0].bold = True
    set_table_geometry(table, [3000, 1700, 4660])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(72)
add_title("RELATÓRIO DA SPRINT", 29, NAVY, True, WD_ALIGN_PARAGRAPH.CENTER, 8)
add_title("Grow Finance Hub", 17, BLUE, False, WD_ALIGN_PARAGRAPH.CENTER, 28)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(5)
r = p.add_run("Período consolidado até 06 de agosto de 2026")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor.from_string(MUTED)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Obrigações  |  Documentos e robô  |  Calendário operacional")
r.font.size = Pt(10.5)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(BLUE)
doc.add_paragraph().paragraph_format.space_after = Pt(150)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Documento interno de acompanhamento técnico e operacional")
r.font.size = Pt(9.5)
r.font.italic = True
r.font.color.rgb = RGBColor.from_string(MUTED)
doc.add_page_break()

doc.add_heading("1. Resumo executivo", level=1)
p = doc.add_paragraph(style="Lead")
p.add_run("A sprint avançou em três frentes principais: regras e geração das obrigações, documentos e reconhecimento pelo robô e calendário operacional.")
doc.add_paragraph("O fluxo de obrigações está mais centralizado, auditável e seguro. O robô deixou de depender do nome do arquivo, passou a trabalhar com modelos validados e correção manual rastreável. O calendário foi transformado em uma visão operacional agrupada por vencimento, obrigação e cliente.")
doc.add_paragraph("Estado geral: núcleo funcional implementado e validações automatizadas aprovadas. Permanecem como etapas de encerramento a homologação visual com usuário autenticado, a validação com documentos reais e a organização do conjunto local em commit e push.")

doc.add_heading("Visão consolidada", level=2)
add_status_table([
    ("Motor de geração", "Implementado", "Geração mensal centralizada no dia 25."),
    ("Regras de competência", "Implementado", "Cenários principais cobertos; requer homologação ampla."),
    ("Documentos e robô", "Implementado", "Modelo seguro ativo, aguardando amostras reais."),
    ("Lista de entregas", "Implementado", "Fluxo de documento, envio e acesso centralizado."),
    ("Calendário operacional", "Implementado", "Agenda por período, obrigação e cliente."),
    ("Homologação final", "Pendente", "Validar sessão autenticada e dados reais."),
    ("Commit e push", "Pendente", "Alterações atuais ainda estão no working tree."),
])

doc.add_heading("2. Regras e geração de obrigações", level=1)
doc.add_heading("Entregas realizadas", level=2)
add_bullets([
    "Centralização da geração mensal com dia fixo definido como 25 de cada mês.",
    "Tratamento de competência anterior e vigente.",
    "Revisão de vencimentos por dia fixo, dia útil, periodicidade mensal, trimestral, anual e múltiplas datas anuais.",
    "Inclusão de feriados brasileiros para o cálculo de dias úteis.",
    "Deduplicação de competências e tarefas, com preservação do vínculo por regime tributário.",
    "Separação entre instância da obrigação, tarefa operacional e representação no calendário.",
    "Remoção de eventos redundantes do calendário e correção da origem visual das tarefas.",
    "Abertura da tarefa operacional diretamente a partir do calendário.",
])
doc.add_heading("Resultado operacional", level=2)
doc.add_paragraph("A instância permanece como registro central da competência. A tarefa controla a execução e o calendário representa o vencimento. Essa estrutura reduz caminhos paralelos de atualização e facilita auditoria, deduplicação e manutenção da regra de negócio.")

doc.add_heading("3. Documentos e robô", level=1)
doc.add_heading("Reconhecimento seguro", level=2)
add_bullets([
    "O nome do arquivo deixou de participar da identificação do documento.",
    "CNPJ e competência são extraídos do conteúdo do PDF.",
    "As áreas marcadas no documento-modelo orientam a leitura e são obrigatórias para automação.",
    "Modelos possuem versão e estados: rascunho, em validação, aprovado e inativo.",
    "Somente modelos aprovados podem vincular documentos automaticamente.",
    "A aprovação automática exige pelo menos cinco validações, quatro acertos e nenhum falso positivo.",
    "Alterações de áreas ou reprocessamento incrementam a versão e reiniciam a validação.",
])
doc.add_heading("Correção manual e auditoria", level=2)
add_bullets([
    "A correção manual atualiza cliente, obrigação, competência, instância e tarefa operacional.",
    "Toda correção exige justificativa e preserva a leitura original.",
    "A mesma amostra não pode ser contabilizada repetidamente nas métricas do modelo.",
    "Documento validado segue para o andamento da tarefa e a obrigação segue para revisão.",
    "A interface apresenta cobertura dos modelos, evidências, acertos e falsos positivos.",
])
doc.add_heading("Situação atual", level=2)
doc.add_paragraph("Há somente um documento-modelo cadastrado. Ele possui as áreas necessárias, mas permanece em validação. Até atingir a amostragem mínima, os documentos relacionados devem passar por revisão manual. Essa trava reduz o risco de associações incorretas.")

doc.add_heading("4. Central e entrega de documentos", level=1)
add_bullets([
    "Lista de entregas criada como visão simplificada das obrigações.",
    "Filtros por cliente, setor, competência, status e período.",
    "Obrigações expansíveis com protocolo, destinatário, documento e envio.",
    "Envio simplificado a partir do arquivo anexado, mantendo atualização de instância, tarefa e andamento.",
    "Envio ao cliente por link, sem anexar diretamente o PDF.",
    "Registro de acesso com data, meio, cliente e documento.",
    "Lista de entregas replicada na página do cliente.",
    "Separação explícita entre documento esperado e documento efetivamente enviado.",
])

doc.add_heading("5. Calendário operacional", level=1)
doc.add_heading("Visão e priorização", level=2)
add_bullets([
    "Agenda operacional mensal ou do dia, sem exigir a abertura individual de cada data.",
    "Hierarquia por data, obrigação e cliente/competência.",
    "Grupos críticos iniciam abertos e empresas são ordenadas por criticidade.",
    "Indicadores de itens em aberto, atrasados, concluídos e em revisão.",
    "Filtros por período, setor, status, obrigação, cliente, CNPJ e somente atrasadas.",
    "Colaboradores permanecem restritos ao próprio setor.",
    "Instâncias canceladas ou substituídas são excluídas da visão.",
    "Tarefas comuns e eventos permanecem separados das obrigações.",
    "Clique em uma competência abre sua tarefa operacional.",
])
doc.add_heading("Otimização do banco", level=2)
doc.add_paragraph("Foram criados índices específicos para eventos por organização/data/setor, tarefas por organização/vencimento/setor/status e instâncias ativas por organização/vencimento/status/obrigação. As consultas passaram a limitar o conjunto por organização, mês e filtros estruturados.")

doc.add_heading("6. Outras melhorias incluídas", level=1)
add_bullets([
    "Chat interno ampliado e alinhado visualmente ao módulo de WhatsApp.",
    "Separadores de mensagens para hoje, ontem e datas anteriores.",
    "Menu de obrigações reorganizado em Central de documentos, Catálogo e Lista de entregas.",
    "Obrigações e histórico acessíveis na página do cliente.",
    "Integração mais clara entre calendário, tarefa e instância.",
])

doc.add_heading("7. Qualidade e validações", level=1)
add_status_table([
    ("Testes automatizados", "Implementado", "42 arquivos e 137 testes aprovados."),
    ("Análise estática", "Implementado", "ESLint concluído sem erros."),
    ("Build de produção", "Implementado", "Compilação Vite concluída."),
    ("Supabase", "Implementado", "Migrações recentes, função e índices aplicados."),
    ("Verificação de rota", "Implementado", "Sem overlay do Vite; login protege a rota."),
    ("Homologação autenticada", "Pendente", "Requer sessão real e dados operacionais."),
])

doc.add_heading("8. Pontos pendentes e riscos", level=1)
doc.add_heading("Homologação funcional", level=2)
add_bullets([
    "Validar o calendário com usuário autenticado e dados reais.",
    "Confirmar a abertura lateral das tarefas em todos os estados.",
    "Validar responsividade e filtros combinados com grande volume.",
    "Confirmar o cálculo de atraso nas datas reais da operação.",
])
doc.add_heading("Validação do robô", level=2)
add_bullets([
    "Cadastrar mais documentos-modelo e validar pelo menos cinco documentos reais por modelo.",
    "Registrar falsos positivos e correções.",
    "Testar diferentes formatos de PDF e documentos digitalizados com OCR.",
])
doc.add_heading("Organização do código", level=2)
doc.add_paragraph("O diretório possui um volume significativo de alterações ainda não commitadas, incluindo novas migrações, testes, componentes e funções. Alterações anteriores e atuais estão misturadas no mesmo working tree.")

doc.add_page_break()
doc.add_heading("9. Próximos passos recomendados", level=1)
add_numbered([
    "Separar arquivos temporários e resultados de verificação.",
    "Revisar o diff final e confirmar o escopo de cada frente.",
    "Executar homologação autenticada do calendário e da central de documentos.",
    "Validar documentos reais no fluxo do robô.",
    "Criar commits organizados por frente ou um commit consolidado aprovado.",
    "Fazer push e executar a verificação de deploy.",
    "Registrar evidências finais e encerrar formalmente a sprint.",
])

doc.add_heading("10. Conclusão", level=1)
doc.add_paragraph("A sprint consolidou a base operacional das obrigações: geração previsível, tarefa como controle de execução, documentos rastreáveis, reconhecimento protegido por validação e calendário orientado a vencimentos. O risco técnico principal deixou de ser a ausência de estrutura e passou a ser a necessidade de homologação com volume e documentos reais.")
doc.add_paragraph("Com a conclusão das validações pendentes, organização dos commits e publicação final, o conjunto estará pronto para encerramento e acompanhamento por métricas operacionais.")

doc.core_properties.title = "Relatório da Sprint - Grow Finance Hub"
doc.core_properties.subject = "Consolidação técnica e operacional da sprint"
doc.core_properties.author = "Grow Finance Hub"
doc.core_properties.keywords = "sprint, obrigações, documentos, robô, calendário operacional"
doc.save(OUT)
print(OUT)
